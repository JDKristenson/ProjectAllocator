"""Word document parser using python-docx."""

from docx import Document
from pathlib import Path


def parse_docx(file_path: str | Path) -> str:
    """
    Extract text content from a Word document.

    Args:
        file_path: Path to the .docx file

    Returns:
        Extracted text content as a string
    """
    file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(f"Word file not found: {file_path}")

    if not file_path.suffix.lower() == '.docx':
        raise ValueError(f"Not a Word file: {file_path}")

    doc = Document(file_path)

    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                paragraphs.append(" | ".join(row_text))

    return "\n\n".join(paragraphs)
